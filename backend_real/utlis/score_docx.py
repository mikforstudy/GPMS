from docx import Document
import os


def replace_text_keep_style(run, old_text, new_text):
    """替换Run中的文本并保留样式"""
    if old_text in run.text:
        run.text = run.text.replace(old_text, new_text)


def merge_runs(paragraph):
    """合并段落中相同样式的相邻Run（修正样式判断逻辑）"""
    if len(paragraph.runs) <= 1:
        return

    merged_runs = []
    current_text = ""
    # 初始化为第一个Run的字体属性
    current_font = {
        "name": paragraph.runs[0].font.name,
        "size": paragraph.runs[0].font.size,
        "bold": paragraph.runs[0].font.bold,
        "italic": paragraph.runs[0].font.italic,
        "underline": paragraph.runs[0].font.underline,
    }

    for run in paragraph.runs:
        # 提取当前Run的字体属性
        run_font = {
            "name": run.font.name,
            "size": run.font.size,
            "bold": run.font.bold,
            "italic": run.font.italic,
            "underline": run.font.underline,
        }

        # 判断当前Run的字体是否与合并中的字体一致
        if run_font == current_font:
            current_text += run.text
        else:
            merged_runs.append((current_text, current_font))
            current_text = run.text
            current_font = run_font

    merged_runs.append((current_text, current_font))

    # 清空段落并重新添加合并后的Run
    paragraph.clear()
    for text, font in merged_runs:
        new_run = paragraph.add_run(text)
        # 应用字体属性
        new_run.font.name = font["name"]
        new_run.font.size = font["size"]
        new_run.font.bold = font["bold"]
        new_run.font.italic = font["italic"]
        new_run.font.underline = font["underline"]


def process_paragraph(paragraph, replacements):
    """处理单个段落：合并Run后替换文本"""
    merge_runs(paragraph)  # 预处理合并Run
    for run in paragraph.runs:
        for key, value in replacements.items():
            replace_text_keep_style(run, key, value)


def process_table(table, replacements):
    """处理表格中的单元格"""
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                process_paragraph(para, replacements)


def replace_text_in_docx(input_path, output_path, replacements):
    doc = Document(input_path)

    # 处理所有段落
    for para in doc.paragraphs:
        process_paragraph(para, replacements)

    # 处理所有表格
    for table in doc.tables:
        process_table(table, replacements)

    doc.save(output_path)
    print(f"文件已生成：{output_path}")


# 示例调用
if __name__ == "__main__":
    input_file = r"D:\Code\demo\backend_real\template\本科毕业设计成绩评审表.docx"
    output_folder = "..\\output\\1"
    os.makedirs(output_folder, exist_ok=True)
    output_file = os.path.join(output_folder, "本科毕业设计成绩评审表.docx")

    replacements = {
        "{题目}": "基于Python的智能农业系统设计与实现",
        "{学生姓名}": "student1",
        "{content1}": "评语1",
        "{score1}": "80",
        "{content2}": "评语2",
        "{score2}": "90",
        "{content3}": "评语2",
        "{score3}": "100",
        "{score4}": "95",
    }

    replace_text_in_docx(input_file, output_file, replacements)