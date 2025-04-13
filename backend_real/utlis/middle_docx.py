import os.path
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn

def create_middle_docx(title, teacher_name, student_name, student_id, content1, content6, content7, output_path):
    # 创建一个新的文档对象
    document = Document()

    # 修改默认的 "Normal" 样式
    style = document.styles['Normal']
    style.font.name = '宋体'              # 设置英文字体
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')  # 设置东亚文本字体

    # 如果文档中还有其它样式（如标题、表格等），也可同理修改它们：
    heading1 = document.styles['Heading 1']
    heading1.font.name = '宋体'
    heading1.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    heading1.font.size = Pt(24)

    # 添加标题并居中
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = paragraph.add_run('天津农学院2025届本科毕业论文（设计）中期检查表')
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')  # 设置东亚文本字体

    # 添加表格
    table = document.add_table(rows=3, cols=4)
    table.style = 'Table Grid'

    cells = table.rows[0].cells
    cells[0].text = '题    目'
    cells[1].merge(cells[3]).text = title

    cells = table.rows[1].cells
    cells[0].text = '指导教师'
    cells[1].text = teacher_name
    cells[2].text = '职称'
    cells[3].text = '讲师'

    cells = table.rows[2].cells
    cells[0].text = '学生姓名'
    cells[1].text = student_name
    cells[2].text = '学号'
    cells[3].text = student_id

    # 一、设计（论文）的进展情况：（课题工作量是否适中、已完成任务、未完成任务、完成课题存在问题和解决办法）
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(10)  # 设置段前间距
    paragraph_format.space_after = Pt(0)   # 设置段后间距
    run = paragraph.add_run('一、设计（论文）的进展情况：（课题工作量是否适中、已完成任务、未完成任务、完成课题存在问题和解决办法）')
    run.font.size = Pt(15)

    table = document.add_table(rows=10, cols=1)
    table.style = 'Table Grid'

    merged_cell = table.cell(0, 0)
    for i in range(1, 10):
        merged_cell.merge(table.cell(i, 0))
    merged_cell.text = content1


    # 六、指导教师意见
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(10)  # 设置段前间距
    paragraph_format.space_after = Pt(0)   # 设置段后间距
    run = paragraph.add_run('六、指导教师意见')
    run.font.size = Pt(15)

    table = document.add_table(rows=5, cols=1)
    table.style = 'Table Grid'

    merged_cell = table.cell(0, 0)
    for i in range(1, 5):
        merged_cell.merge(table.cell(i, 0))
    merged_cell.text = content6


    # 七、备注
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(10)  # 设置段前间距
    paragraph_format.space_after = Pt(0)   # 设置段后间距
    run = paragraph.add_run('备注')
    run.font.size = Pt(15)

    table = document.add_table(rows=5, cols=1)
    table.style = 'Table Grid'

    merged_cell = table.cell(0, 0)
    for i in range(1, 5):
        merged_cell.merge(table.cell(i, 0))
    merged_cell.text = content7

    # 添加创建日期
    creation_date = datetime.now().strftime('%Y-%m-%d')
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = paragraph.add_run(f'创建日期: {creation_date}\t')  # 添加制表符，调整位置
    date_run.font.size = Pt(12)
    date_run.element.rPr.get_or_add_rFonts().set(qn('w:eastAsia'), '宋体')  # 设置东亚文本字体

    # 保存文档
    document.save(output_path)

# create_middle_docx(
#     title='基于Python的电商用户行为分析与购买预测算法模型设计',
#     teacher_name='教师1',
#     student_name='学生1',
#     student_id='2025001',
#     content1='摘要1',
#     content6='摘要6',
#     content7='备注7',
#     output_path='../output/中期检查表格.docx'
# )